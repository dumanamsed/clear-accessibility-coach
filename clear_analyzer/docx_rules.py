import re
from io import BytesIO
from docx import Document
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from .models import Finding
from .contrast import contrast_ratio, required_ratio, is_near
from .theme import docx_theme_map, resolve_docx_color, docx_cell_fill
from .altq import alt_quality_findings

# Fonts CLEAR's "Easy to Read" calls out as hard to read: decorative/script
# faces, and a few notoriously low-legibility display fonts. Sans-serif body
# fonts (Arial, Verdana, Calibri, Open Sans...) are the recommendation.
_DECORATIVE_FONTS = {
    "brush script mt", "comic sans ms", "papyrus", "lobster", "pacifico",
    "vivaldi", "edwardian script itc", "monotype corsiva", "freestyle script",
    "bradley hand itc", "chiller", "jokerman", "magneto", "curlz mt",
    "harrington", "mistral", "kunstler script", "blackadder itc", "vladimir script",
    "french script mt", "rage italic", "script mt bold", "segoe script",
}

_IMAGE_EXT_RE = re.compile(
    r"^.+\.(png|jpg|jpeg|gif|bmp|tiff|tif|svg|webp|ico|emf|wmf)$",
    re.IGNORECASE,
)
_WEAK_ALT_PATTERNS = [
    re.compile(r"^image\s*\d*$", re.IGNORECASE),
    re.compile(r"^picture\s*\d*$", re.IGNORECASE),
    re.compile(r"^photo\s*\d*$", re.IGNORECASE),
    re.compile(r"^img[_\s]?\d*$", re.IGNORECASE),
    re.compile(r"^screenshot", re.IGNORECASE),
    re.compile(r"^graphic\s*\d*$", re.IGNORECASE),
    re.compile(r"^chart\s*\d*$", re.IGNORECASE),
]


def analyze_docx(file_bytes: bytes) -> list[Finding]:
    findings = []
    doc = Document(BytesIO(file_bytes))

    _check_images(doc, findings)
    _check_headings(doc, findings)
    _check_hyperlinks(doc, findings)
    _check_paragraphs(doc, findings)
    _check_tables(doc, findings)
    theme_map = docx_theme_map(doc)
    _check_small_fonts(doc, findings)
    _check_all_caps(doc, findings)
    _check_text_contrast(doc, findings, theme_map)
    _check_table_contrast(doc, findings, theme_map)
    _check_fonts(doc, findings)
    _check_justified(doc, findings)

    return findings


from docx.oxml.ns import qn as _qn


def _run_color(run, theme_map):
    """Effective (r,g,b) of a run's text color, resolving theme colors + tints,
    or None for automatic/black. Reads the <w:color> element directly so we see
    theme references python-docx hides behind 'automatic'."""
    rPr = run._element.find(_qn("w:rPr"))
    if rPr is None:
        return None
    color = rPr.find(_qn("w:color"))
    if color is None:
        return None
    return resolve_docx_color(color, theme_map)


def _flag_contrast(rgb, bg, run, location, findings, bg_label=""):
    if rgb is None or bg is None:
        return False
    # Treat text that already contrasts like normal dark-on-light as fine.
    size_pt = run.font.size.pt if run.font.size else None
    font_px = size_pt * 1.333 if size_pt else None
    ratio = contrast_ratio(rgb, bg)
    needed = required_ratio(font_px, bool(run.font.bold))
    if ratio + 0.05 < needed:
        where = bg_label or "a white page"
        findings.append(Finding(
            strand="E",
            severity="warning",
            location=location,
            issue=f"Text color #%02X%02X%02X has only {ratio:.1f}:1 contrast against {where}, "
                  f"below the WCAG 2.2 AA minimum of {needed:g}:1 (Success Criterion 1.4.3). "
                  f"Light or low-contrast text is hard to read." % rgb,
            evidence=run.text.strip()[:80],
        ))
        return True
    return False


def _check_text_contrast(doc, findings, theme_map, max_flags=5):
    """WCAG 1.4.3 on body text. Resolves theme colors (incl. light tints) so
    'light gray' text — the most common real-world failure — is caught. Page
    background assumed white (the Word default). Near-black is skipped so normal
    text never false-positives."""
    WHITE = (255, 255, 255)
    flagged = 0
    for i, para in enumerate(doc.paragraphs, 1):
        if flagged >= max_flags:
            break
        for run in para.runs:
            if not run.text.strip():
                continue
            rgb = _run_color(run, theme_map)
            if rgb is None or is_near(rgb, (0, 0, 0), tol=40):
                continue
            if _flag_contrast(rgb, WHITE, run, f"Paragraph {i}", findings):
                flagged += 1
                break


def _check_table_contrast(doc, findings, theme_map, max_flags=4):
    """WCAG 1.4.3 inside tables: light cell shading behind text, or light text
    on a cell fill. Resolves both the text color AND the cell's shading fill."""
    WHITE = (255, 255, 255)
    flagged = 0
    for ti, table in enumerate(doc.tables, 1):
        if flagged >= max_flags:
            break
        for ri, row in enumerate(table.rows, 1):
            for ci, cell in enumerate(row.cells, 1):
                fill = docx_cell_fill(cell._tc, theme_map)
                bg = fill or WHITE
                for para in cell.paragraphs:
                    for run in para.runs:
                        if not run.text.strip():
                            continue
                        rgb = _run_color(run, theme_map)
                        # If text is default/black, only a dark fill is a problem.
                        if rgb is None or is_near(rgb, (0, 0, 0), tol=40):
                            rgb = (0, 0, 0)
                        loc = f"Table {ti}, row {ri}, col {ci}"
                        label = ("the cell fill #%02X%02X%02X" % fill) if fill else "a white page"
                        if _flag_contrast(rgb, bg, run, loc, findings, label):
                            flagged += 1
                            break
                    if flagged >= max_flags:
                        break
                if flagged >= max_flags:
                    break


def _check_fonts(doc, findings, max_flags=3):
    """CLEAR Easy to Read: 'Use sans-serif fonts such as Arial, Verdana, or Open
    Sans... easier to read than decorative or script fonts.' Flags decorative /
    script faces used for real text."""
    flagged = 0
    seen = set()
    for i, para in enumerate(doc.paragraphs, 1):
        if flagged >= max_flags:
            break
        for run in para.runs:
            name = (run.font.name or "")
            key = name.lower().strip()
            if key in _DECORATIVE_FONTS and len(run.text.strip()) >= 3 and key not in seen:
                seen.add(key)
                findings.append(Finding(
                    strand="E",
                    severity="tip",
                    location=f"Paragraph {i}",
                    issue=f"The font \"{name}\" is decorative/script, which is hard to read in body "
                          f"text. CLEAR recommends clean sans-serif fonts such as Arial, Verdana, "
                          f"or Calibri.",
                    evidence=run.text.strip()[:80],
                ))
                flagged += 1
                break


def _check_justified(doc, findings, max_flags=3):
    """CLEAR Easy to Read: prefer left-aligned over justified text — justification
    creates uneven word spacing that is harder to read, especially for dyslexic
    readers."""
    flagged = 0
    for i, para in enumerate(doc.paragraphs, 1):
        if flagged >= max_flags:
            break
        if para.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY and len(para.text.split()) > 12:
            findings.append(Finding(
                strand="E",
                severity="tip",
                location=f"Paragraph {i}",
                issue="Paragraph is fully justified, which creates uneven word spacing that is "
                      "harder to read. CLEAR recommends left-aligned text.",
                evidence=para.text.strip()[:80],
            ))
            flagged += 1


def _check_small_fonts(doc, findings, max_flags=5):
    """Flag runs with an explicitly tiny font size (< 10pt)."""
    flagged = 0
    for i, para in enumerate(doc.paragraphs, 1):
        if flagged >= max_flags:
            break
        for run in para.runs:
            size = run.font.size
            if size is not None and size.pt < 10 and run.text.strip():
                findings.append(Finding(
                    strand="E",
                    severity="tip",
                    location=f"Paragraph {i}",
                    issue=f"Text uses a very small font size ({size.pt:.0f}pt). Consider 11pt or larger for body text.",
                    evidence=run.text.strip()[:80],
                ))
                flagged += 1
                break


def _check_all_caps(doc, findings, max_flags=3):
    """Flag long stretches of ALL-CAPS text, which are harder to read and may
    be spelled out letter-by-letter by some screen readers."""
    flagged = 0
    for i, para in enumerate(doc.paragraphs, 1):
        if flagged >= max_flags:
            break
        text = para.text.strip()
        letters = [c for c in text if c.isalpha()]
        if len(letters) > 30 and all(c.isupper() for c in letters):
            findings.append(Finding(
                strand="E",
                severity="tip",
                location=f"Paragraph {i}",
                issue="Long stretch of ALL-CAPS text. Mixed case is easier to read; use bold or a heading style for emphasis instead.",
                evidence=text[:80],
            ))
            flagged += 1


def _find_images_in_element(element):
    """Find all drawing/image elements in an XML element tree.

    Images in DOCX can appear as:
    - w:drawing > wp:inline > a:graphic > a:graphicData > pic:pic
    - w:drawing > wp:anchor > a:graphic > a:graphicData > pic:pic
    - w:pict > v:shape > v:imagedata (legacy VML images)

    Alt text lives in:
    - wp:inline/wp:anchor > wp:docPr[@descr]
    - w:pict > v:shape[@alt]
    """
    images = []

    # Modern DrawingML images: wp:inline and wp:anchor
    for drawing in element.iter(qn("w:drawing")):
        # Check wp:inline elements
        for inline in drawing.iter(qn("wp:inline")):
            docPr = inline.find(qn("wp:docPr"))
            if docPr is not None:
                descr = docPr.get("descr", "")
                name = docPr.get("name", "Image")
                # Check for decorative flag
                decorative = False
                for child in docPr:
                    if "decorative" in child.tag.lower():
                        if child.get("val", "") == "1":
                            decorative = True
                images.append({
                    "name": name,
                    "alt_text": descr.strip() if descr else "",
                    "decorative": decorative,
                    "type": "inline",
                })

        # Check wp:anchor elements (floating images)
        for anchor in drawing.iter(qn("wp:anchor")):
            docPr = anchor.find(qn("wp:docPr"))
            if docPr is not None:
                descr = docPr.get("descr", "")
                name = docPr.get("name", "Image")
                decorative = False
                for child in docPr:
                    if "decorative" in child.tag.lower():
                        if child.get("val", "") == "1":
                            decorative = True
                images.append({
                    "name": name,
                    "alt_text": descr.strip() if descr else "",
                    "decorative": decorative,
                    "type": "anchor",
                })

    # Legacy VML images: w:pict > v:shape
    VML_NS = "urn:schemas-microsoft-com:vml"
    for pict in element.iter(qn("w:pict")):
        for vshape in pict.iter(f"{{{VML_NS}}}shape"):
            alt = vshape.get("alt", "")
            # Check for imagedata child to confirm it's actually an image
            has_image = False
            for child in vshape.iter(f"{{{VML_NS}}}imagedata"):
                has_image = True
                break
            if has_image:
                images.append({
                    "name": vshape.get("id", "VML Image"),
                    "alt_text": alt.strip() if alt else "",
                    "decorative": False,
                    "type": "vml",
                })

    return images


def _evaluate_alt_text(img, location, findings):
    """Evaluate an image's alt text and add findings if problematic."""
    if img["decorative"]:
        return
    alt = img["alt_text"]
    name = img["name"]

    if not alt:
        findings.append(Finding(
            strand="A",
            severity="critical",
            location=location,
            issue="Image is missing alt text.",
            evidence=f"Image: {name}",
        ))
    elif _IMAGE_EXT_RE.match(alt):
        findings.append(Finding(
            strand="A",
            severity="critical",
            location=location,
            issue="Image alt text is a filename, not a meaningful description.",
            evidence=f"Image: {name}, alt text: \"{alt}\"",
        ))
    elif any(p.match(alt) for p in _WEAK_ALT_PATTERNS):
        findings.append(Finding(
            strand="A",
            severity="warning",
            location=location,
            issue="Image alt text appears to be generic or auto-generated. Consider a description of what the image conveys.",
            evidence=f"Image: {name}, alt text: \"{alt}\"",
        ))
    else:
        findings.extend(alt_quality_findings(alt, location))


def _check_images(doc, findings):
    """Check all images in the document body, headers, and footers."""
    image_count = 0

    # Check main document body paragraphs
    for para_idx, para in enumerate(doc.paragraphs, 1):
        images = _find_images_in_element(para._element)
        for img in images:
            image_count += 1
            _evaluate_alt_text(img, f"Paragraph {para_idx}", findings)

    # Check images inside tables
    for tbl_idx, table in enumerate(doc.tables, 1):
        for row_idx, row in enumerate(table.rows, 1):
            for cell_idx, cell in enumerate(row.cells, 1):
                images = _find_images_in_element(cell._element)
                for img in images:
                    image_count += 1
                    _evaluate_alt_text(img, f"Table {tbl_idx}, Row {row_idx}, Cell {cell_idx}", findings)

    # Check headers and footers
    for section in doc.sections:
        for header_footer_name, hf in [
            ("Header", section.header),
            ("Footer", section.footer),
        ]:
            if hf and hf.is_linked_to_previous is False or (hf and hasattr(hf, '_element')):
                try:
                    for para in hf.paragraphs:
                        images = _find_images_in_element(para._element)
                        for img in images:
                            image_count += 1
                            _evaluate_alt_text(img, header_footer_name, findings)
                except Exception:
                    pass


def _check_headings(doc, findings):
    heading_styles = []
    for i, para in enumerate(doc.paragraphs, 1):
        style_name = para.style.name if para.style else ""
        if style_name.startswith("Heading"):
            try:
                level = int(style_name.replace("Heading ", "").strip())
                heading_styles.append((i, level))
            except ValueError:
                heading_styles.append((i, 0))

    if not heading_styles:
        findings.append(Finding(
            strand="L",
            severity="critical",
            location="Entire document",
            issue="No heading styles are used in this document. Only the Normal style was detected.",
            evidence="All paragraphs use Normal or non-heading styles.",
        ))
        return

    for j in range(1, len(heading_styles)):
        prev_para, prev_level = heading_styles[j - 1]
        curr_para, curr_level = heading_styles[j]
        if curr_level > prev_level + 1:
            findings.append(Finding(
                strand="L",
                severity="warning",
                location=f"Paragraph {curr_para}",
                issue=f"Heading level skips from H{prev_level} to H{curr_level}.",
                evidence=f"Expected H{prev_level + 1} but found H{curr_level}.",
            ))


def _check_hyperlinks(doc, findings):
    generic_texts = {"click here", "here", "link", "read more", "learn more"}
    for i, para in enumerate(doc.paragraphs, 1):
        hyperlinks = para._element.findall(f".//{qn('w:hyperlink')}")
        for hl in hyperlinks:
            text_nodes = hl.findall(f".//{qn('w:t')}")
            link_text = "".join(t.text for t in text_nodes if t.text).strip()
            if link_text.lower() in generic_texts:
                findings.append(Finding(
                    strand="E",
                    severity="tip",
                    location=f"Paragraph {i}",
                    issue="Hyperlink uses generic text that does not describe its destination.",
                    evidence=f'Link text: "{link_text}"',
                ))


def _check_paragraphs(doc, findings):
    for i, para in enumerate(doc.paragraphs, 1):
        text = para.text.strip()
        if not text:
            continue
        word_count = len(text.split())
        if word_count > 150:
            findings.append(Finding(
                strand="E",
                severity="tip",
                location=f"Paragraph {i}",
                issue=f"Long paragraph ({word_count} words). Consider breaking it into shorter chunks.",
                evidence=text[:120] + "...",
            ))


def _check_tables(doc, findings):
    for i, table in enumerate(doc.tables, 1):
        if not table.rows:
            continue
        first_row = table.rows[0]
        has_header = False
        row_element = first_row._tr
        trPr = row_element.find(qn("w:trPr"))
        if trPr is not None:
            tblHeader = trPr.find(qn("w:tblHeader"))
            if tblHeader is not None:
                has_header = True
        if not has_header:
            findings.append(Finding(
                strand="L",
                severity="warning",
                location=f"Table {i}",
                issue="Table does not have a designated header row.",
                evidence="First row is not marked as a header, which impacts screen reader navigation.",
            ))
